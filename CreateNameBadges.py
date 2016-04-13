import openpyxl
import docx
import os
import time


class Attendee():

    def __init__(self, last_name, first_name, corporation, city, country):
        self.last_name = last_name
        self.first_name = first_name
        self.corporation = corporation
        self.city = city
        self.country = country

os.chdir('/Users/DoctorSushi/Downloads')
saveFolder = os.path.join(os.getcwd(), time.strftime("%m-%d-%Y") + ' Name Badges')
if not os.path.isdir(saveFolder):
    os.mkdir(saveFolder)
    
doc = docx.Document('SEA2016 Badge Template Attendee.docx')
#workbook = openpyxl.load_workbook('badgetemplate.xlsx')
workbook = openpyxl.load_workbook(input('Enter file name for attendee list:\n'))
sheet = workbook.get_sheet_by_name('Sheet1')
attendees = []
for i in range(2, len(sheet.rows)+ 1):
    attendees.append(Attendee(sheet['A'+ str(i)].value, sheet['B'+ str(i)].value, sheet['F'+ str(i)].value, sheet['I'+ str(i)].value, sheet['L'+ str(i)].value))   

for i in range(len(attendees)):
    if (attendees[i].last_name == None or attendees[i].first_name == None):
        continue
    
    for j in range(int(len(doc.paragraphs))):
        if doc.paragraphs[j].text == 'Name':
            full_name = attendees[i].first_name + ' ' + attendees[i].last_name
            doc.paragraphs[j].runs[0].text = full_name
            
        if doc.paragraphs[j].text == 'Corporation':
            doc.paragraphs[j].runs[0].text = attendees[i].corporation
            
        if doc.paragraphs[j].text == 'City, Country':
            full_location = attendees[i].city + ', ' + attendees[i].country
            doc.paragraphs[j].runs[0].text = full_location

    if full_name != '':
        filePath = os.path.join(saveFolder, full_name + ' ' + attendees[i].corporation + '.docx')
        print(filePath)
        doc.save(filePath)
    doc = docx.Document('SEA2016 Badge Template Attendee.docx')
#update workbook with 'badge setup column'

